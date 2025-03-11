import json
import requests
import docx

from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm

from time import sleep
from bs4 import BeautifulSoup
from lxml import etree

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException

import tkinter as tk
from tkinter import ttk

def get_max_page(url: str) -> int:
    s = requests.Session()
    headers = {
        "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
    }
    r = s.get(url, headers=headers)
    soup = BeautifulSoup(r.content, "html.parser")
    dom = etree.HTML(str(soup))
    pages = dom.xpath(
        '//*[@id="__next"]/div/main/div/div[2]/div/div/div/div[4]/div'
    )
    try:
        pages = pages[0].findall(".//a")
        return int(pages[-1].text)
    except IndexError:
        return int(1)

def get_links_from_page(url: str) -> list[str]:
    s = requests.Session()
    headers = {
        "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
    }
    r = s.get(url, headers=headers)
    soup = BeautifulSoup(r.content, "html.parser")
    dom = etree.HTML(str(soup))
    cards = dom.xpath(
        '//*[@id="__next"]/div/main/div/div[2]/div/div/div/div[3]/div/div'
    )
    cards = cards[0].findall(".//a")
    links = [f'https://www.culture.ru{card.get("href")}' for card in cards]
    return links


def parse_data(url: str) -> list[dict]:
    text.insert(1.0, "[+] Осуществляется парсинг данных. Это может занять некоторое время..." + "\n")
    max_page: int = get_max_page(url)
    urls = url
    all_links: list[str] = []

    for page in range(1, max_page + 1):
        url = f"{urls}?page={page}"
        links = get_links_from_page(url)
        all_links.extend(links)

    text.insert(1.0, f"[+] Найдено {len(all_links)} событий." + "\n")

    options = webdriver.EdgeOptions()
    driver = webdriver.Edge(options=options)

    results = []
    for link in all_links:
        text.insert(1.0, f"[+] Обрабатывается событие: {link}" + "\n")
        driver.get(link)
        page_content = driver.page_source

        soup = BeautifulSoup(page_content, "html.parser")

        buy_tickets_link = ""

        if soup.find("div", {"class": "Rv1DR Rh1E3 afWIX"}):
            bt_checker = driver.find_element(
                By.CSS_SELECTOR, '#__next > div > div.Rv1DR.Rh1E3.afWIX > div._0EYFN > div.HPr50 > button')
            sleep(1)
            bt_checker.click()

        wait = WebDriverWait(driver, 5)
        button = wait.until(
            EC.visibility_of_element_located((By.CLASS_NAME, "doHY5")))

        actions = ActionChains(driver)
        actions.move_to_element(button).perform()

        button.click()

        sleep(0.3)
        try:
            buy_tickets_link = driver.find_element(
                By.XPATH,
                "/html/body/div[7]/div/div/div/div[3]/a"
            ).get_attribute("href")
        except NoSuchElementException:
            sleep(2)
            buy_tickets_link = driver.find_element(
                By.XPATH,
                "/html/body/div[7]/div/div/div/div[3]/a"
            ).get_attribute("href")
            print("trouble with link")
        dom = etree.HTML(str(soup))

        name = dom.xpath(
            '//*[@id="__next"]/div/main/div/div[3]/div[1]/div/div/h1/text()'
        )[0]

        cur_data_json = {
            "name": name,
            "link": link,
            "place": soup.find("div", {"class": "uMrgA"}).text,
            "date": soup.find("div", {"class": "Jds71"}).find_all("div", {"class": "_19IwE"})[0].text,
            "age": soup.find("div", {"class": "Jds71"}).find_all("div", {"class": "_19IwE"})[1].text,
            "text": soup.find("div", {"class": "xZmPc"}).text,
            "address": soup.find("div", {"class": "SHIlp"}).text,
            "price": soup.find("div", {"class": "O7bBt"}).text,
            "time": soup.find("div", {"class": "v5z9s"}).text,
            "buy_tickets_link": buy_tickets_link
        }
        results.append(cur_data_json)

    driver.close()
    text.insert(1.0, "[+] Парсинг завершен." + "\n")
    return results

def make_third_column_wider(doc):
    table = doc.tables[0]

    # Установка ширины столбцов
    table.columns[0].width = Cm(3)  # Ширина первого столбца
    table.columns[1].width = Cm(5)  # Ширина второго столбца
    table.columns[2].width = Cm(8)  # Ширина третьего столбца
    table.columns[3].width = Cm(4)  # Ширина четвертого столбца

def write_data_to_docx(data: list[dict]) -> None:
    text.insert(1.0, "[+] Запись данных в docx-файл..." + "\n")
    sleep(1)
    doc = docx.Document()
    # Добавление параграфа
    paragraph = doc.add_paragraph("Афиша программы «Пушкинская карта»")

    # Изменение шрифта
    run = paragraph.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Изменение выравнивания
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=len(data) + 1, cols=4)

    table.cell(0, 0).text = "Дата/время"
    table.cell(0, 1).text = "Наименование мероприятия/Ссылка на мероприятие"
    table.cell(0, 2).text = "Краткое описание"
    table.cell(0, 3).text = "Купить билет"

    for i, item in enumerate(data, 1):
        table.cell(i, 0).text = f'{item["date"]}/{item["time"]}'
        table.cell(i, 1).text = f'{item["name"]} - {item["link"]}'
        table.cell(i, 2).text = item["text"]
        table.cell(i, 3).text = item["buy_tickets_link"]

    # Изменение шрифта для данных в таблице
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    # Сделать подпись к данным в жирном стиле
    for cell in table.row_cells(0):
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.bold = True

    make_third_column_wider(doc)

    doc.save("results.docx")
    text.insert(1.0, "[+] Данные записаны в docx-файл results.docx" + "\n")

urls = {
    "Абакан": "https://www.culture.ru/afisha/respublika-hakasiya-abakan/pushkinskaya-karta",
    "Черногорск": "https://www.culture.ru/afisha/respublika-hakasiya-chernogorsk/pushkinskaya-karta",
    "Сорск": "https://www.culture.ru/afisha/respublika-hakasiya-sorsk/pushkinskaya-karta",
    "Саяногорск": "https://www.culture.ru/afisha/respublika-hakasiya-sayanogorsk/pushkinskaya-karta",
    "Абаза": "https://www.culture.ru/afisha/respublika-hakasiya-abaza/pushkinskaya-karta",
    "Усть-Абаканский район": "https://www.culture.ru/afisha/respublika-hakasiya-ust-abakanskiy-rayon/pushkinskaya-karta",
    "Таштыпский район": "https://www.culture.ru/afisha/respublika-hakasiya-tashtypskiy-rayon/pushkinskaya-karta",
    "Аскизский район": "https://www.culture.ru/afisha/respublika-hakasiya-askizskiy-rayon/pushkinskaya-karta",
    "Алтайский район": "https://www.culture.ru/afisha/respublika-hakasiya-altaiskii-raion/pushkinskaya-karta",
    "Орджоникидзевский район": "https://www.culture.ru/afisha/respublika-hakasiya-ordzhonikidzevskiy-rayon/pushkinskaya-karta",
    "Ширинский район": "https://www.culture.ru/afisha/respublika-hakasiya-shirinskiy-rayon/pushkinskaya-karta",
}

def start():
    i = combobox.get()
    url = urls[i]
    results = parse_data(url)

    with open("results.json", "w", encoding="utf-8") as file:
        json.dump(results, file, indent=4, ensure_ascii=False)
    text.insert(1.0, "[+] Результаты сохранены в results.json" + "\n")

    with open("results.json", "r", encoding="utf-8") as file:
        data = json.load(file)
    text.insert(1.0, "[+] Результаты сохранены в results.docx" + "\n")
    write_data_to_docx(data)

# GUI
# Главное окно
root = tk.Tk()
root.title('Парсер')

inputFrame = tk.Frame(root)
inputFrame.grid(row=0, column=0, ipadx=6, ipady=6, padx=4, pady=4)
sity = ["Абакан", "Черногорск", "Сорск", "Саяногорск", "Абаза", "Усть-Абаканский район",
        "Таштыпский район", "Аскизский район", "Алтайский район",
        "Орджоникидзевский район", "Ширинский район"]
sity_var = tk.StringVar(value=sity[0])
combobox = ttk.Combobox(inputFrame, textvariable=sity_var, values=sity)
combobox.grid(row=0, column=0, columnspan=2)

btn = tk.Button(inputFrame, text="Старт", command=start)
btn.grid(row=1, column=0, ipadx=50, ipady=6, padx=5, pady=5)

outputFrame = tk.Frame(root)
outputFrame.grid(row=0, column=1, columnspan=2)
text = tk.Text(outputFrame, height=20, width=60, font='Arial 14', wrap=tk.WORD)
text.pack()

root.mainloop()